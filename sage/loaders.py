"""Load text from supported documentation file types (local-only processing)."""

from __future__ import annotations

import csv
import json
import os
import re
from io import StringIO
from pathlib import Path

from sage.config import (
    ENV_FILE_NAMES,
    SENSITIVE_EXTENSIONS,
    SKIP_DIR_NAMES,
    SKIP_FILE_NAMES,
    SKIP_FILE_SUBSTRINGS,
    SUPPORTED_EXTENSIONS,
)

_SKIP_LOWER = {n.lower() for n in SKIP_DIR_NAMES}
_SKIP_FILES_LOWER = {n.lower() for n in SKIP_FILE_NAMES}
_SKIP_SUBSTR_LOWER = tuple(s.lower() for s in SKIP_FILE_SUBSTRINGS)


def is_env_file(path: Path) -> bool:
    """True for .env, .env.local, app.env, etc."""
    name = path.name.lower()
    if name in ENV_FILE_NAMES:
        return True
    if name.startswith(".env."):
        return True
    if path.suffix.lower() == ".env":
        return True
    return False


def is_supported_file(path: Path) -> bool:
    if not path.is_file():
        return False
    name_l = path.name.lower()
    if name_l in _SKIP_FILES_LOWER:
        return False
    if any(s in name_l for s in _SKIP_SUBSTR_LOWER):
        return False
    if is_env_file(path):
        return True
    return path.suffix.lower() in SUPPORTED_EXTENSIONS


def is_sensitive_source(path: Path) -> bool:
    """Files that may contain secrets; still local-only in Project Sage."""
    if is_env_file(path):
        return True
    return path.suffix.lower() in SENSITIVE_EXTENSIONS


def iter_supported_files(root: Path) -> list[Path]:
    """Walk a folder and return supported files, skipping noisy directories."""
    root = root.resolve()
    found: list[Path] = []
    if not root.is_dir():
        return found

    for dirpath, dirnames, filenames in os.walk(root):
        # Prune skip dirs and hidden folders in-place
        dirnames[:] = [
            d
            for d in dirnames
            if d.lower() not in _SKIP_LOWER and not d.startswith(".")
        ]
        for name in filenames:
            p = Path(dirpath) / name
            if is_supported_file(p):
                found.append(p)
    return sorted(found)


def load_file_text(path: Path) -> str:
    """Extract plain text from a supported file. Never uploads off-machine."""
    suffix = path.suffix.lower()
    if is_env_file(path):
        return _load_env(path)
    if suffix in {".txt", ".md", ".markdown", ".log"}:
        return _read_text(path)
    if suffix == ".pdf":
        return _load_pdf(path)
    if suffix == ".csv":
        return _load_csv(path)
    if suffix == ".json":
        return _load_json(path)
    if suffix == ".jsonl":
        return _load_jsonl(path)
    if suffix in {".yaml", ".yml"}:
        return _load_yaml(path)
    if suffix == ".docx":
        return _load_docx(path)
    if suffix == ".doc":
        return _load_doc_legacy(path)
    raise ValueError(f"Unsupported file type: {suffix or path.name}")


def _read_text(path: Path) -> str:
    raw = path.read_bytes()
    for enc in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _load_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts: list[str] = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            parts.append(f"--- Page {i + 1} ---\n{text}")
    return "\n\n".join(parts)


def _load_csv(path: Path) -> str:
    text = _read_text(path)
    reader = csv.reader(StringIO(text))
    rows = list(reader)
    if not rows:
        return ""
    lines = [" | ".join(cell.strip() for cell in row) for row in rows]
    return "\n".join(lines)


def _load_json(path: Path) -> str:
    text = _read_text(path)
    try:
        data = json.loads(text)
        return json.dumps(data, indent=2, ensure_ascii=False)
    except json.JSONDecodeError:
        return text


def _format_jsonl_record(obj: object, index: int) -> str:
    """Turn one NDJSON object into searchable text (coding-notes aware)."""
    if not isinstance(obj, dict):
        return json.dumps(obj, ensure_ascii=False, default=str)

    # log-sage coding-notes.jsonl shape
    coding_keys = {
        "session_id",
        "session_title",
        "user_query",
        "kind",
        "project",
        "message",
        "preview",
        "timestamp",
    }
    if coding_keys & set(obj.keys()):
        lines: list[str] = [f"### Coding note #{index + 1}"]
        for key, label in (
            ("timestamp", "When"),
            ("project", "Project"),
            ("kind", "Kind"),
            ("session_title", "Session"),
            ("session_id", "Session ID"),
            ("user_query", "User request"),
            ("message", "Notes / instructions"),
            ("preview", "Preview"),
        ):
            val = obj.get(key)
            if val is None or val == "":
                continue
            text = str(val).strip()
            if not text:
                continue
            lines.append(f"{label}: {text}")
        # Any remaining useful fields
        for key, val in obj.items():
            if key in {
                "timestamp",
                "project",
                "kind",
                "session_title",
                "session_id",
                "user_query",
                "message",
                "preview",
                "level",
                "logger",
            }:
                continue
            if val is None or val == "":
                continue
            if isinstance(val, (dict, list)):
                lines.append(
                    f"{key}: {json.dumps(val, ensure_ascii=False, default=str)[:2000]}"
                )
            else:
                lines.append(f"{key}: {val}")
        return "\n".join(lines)

    return json.dumps(obj, indent=2, ensure_ascii=False, default=str)


def _load_jsonl(path: Path) -> str:
    """
    Load newline-delimited JSON (NDJSON / .jsonl).

    Optimized for log-sage ``coding-notes.jsonl``: one record per line becomes
    a labeled block so chunking can keep session notes coherent for RAG.
    """
    parts: list[str] = [f"# JSONL source: {path.name}"]
    try:
        raw_lines = path.read_text(encoding="utf-8", errors="replace").splitlines()
    except OSError as e:
        return f"[Could not read jsonl {path.name}: {e}]"

    n_ok = 0
    n_bad = 0
    for i, line in enumerate(raw_lines):
        stripped = line.strip()
        if not stripped:
            continue
        try:
            obj = json.loads(stripped)
        except json.JSONDecodeError:
            n_bad += 1
            parts.append(f"### Line {i + 1} (raw)\n{stripped[:4000]}")
            continue
        n_ok += 1
        parts.append(_format_jsonl_record(obj, n_ok - 1))

    parts.insert(1, f"# Records parsed: {n_ok}" + (f" | bad lines: {n_bad}" if n_bad else ""))
    # Blank line between records helps chunk boundary detection
    return "\n\n".join(parts)


def _load_yaml(path: Path) -> str:
    """Parse YAML safely and re-serialize for stable searchable text."""
    text = _read_text(path)
    try:
        import yaml  # PyYAML
    except ImportError:
        return (
            f"[YAML file {path.name}: PyYAML not installed; indexing raw text]\n{text}"
        )
    try:
        data = yaml.safe_load(text)
    except yaml.YAMLError as e:
        return f"[YAML parse note for {path.name}: {e}]\n{text}"
    if data is None:
        return ""
    try:
        return yaml.safe_dump(
            data,
            default_flow_style=False,
            allow_unicode=True,
            sort_keys=False,
        )
    except yaml.YAMLError:
        return text


_ENV_LINE = re.compile(
    r"^\s*(?:export\s+)?([A-Za-z_][A-Za-z0-9_]*)\s*=\s*(.*)$"
)


def _load_env(path: Path) -> str:
    """
    Load dotenv-style files as searchable text (local only).

    Keeps key names and values for local RAG. Data is embedded into the
    on-disk Chroma store under data/ (gitignored) and queried via local Ollama —
    nothing is sent to external cloud APIs by Project Sage.
    """
    text = _read_text(path)
    lines_out: list[str] = [
        f"# Local env file: {path.name}",
        "# Stored only in local Project Sage index (data/). Not uploaded externally.",
    ]
    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped or stripped.startswith("#"):
            if stripped.startswith("#"):
                lines_out.append(stripped)
            continue
        m = _ENV_LINE.match(line)
        if not m:
            lines_out.append(line)
            continue
        key, value = m.group(1), m.group(2).strip()
        # Strip optional surrounding quotes for cleaner chunks
        if len(value) >= 2 and value[0] == value[-1] and value[0] in "\"'":
            value = value[1:-1]
        lines_out.append(f"{key}={value}")
    return "\n".join(lines_out)


def _load_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    paras = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                paras.append(" | ".join(cells))
    return "\n\n".join(paras)


def _load_doc_legacy(path: Path) -> str:
    """Best-effort for old .doc binary format."""
    try:
        raw = path.read_bytes()
        chunks: list[str] = []
        current: list[str] = []
        for b in raw:
            if 32 <= b < 127 or b in (9, 10, 13):
                current.append(chr(b))
            else:
                if len(current) >= 40:
                    chunks.append("".join(current))
                current = []
        if len(current) >= 40:
            chunks.append("".join(current))
        text = "\n".join(c.strip() for c in chunks if c.strip())
        if len(text) < 50:
            return (
                f"[Limited extraction from legacy .doc: {path.name}. "
                "Prefer converting to .docx or .pdf for better results.]"
            )
        return text
    except OSError as e:
        return f"[Could not read .doc file {path.name}: {e}]"
